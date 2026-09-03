"""Read text out of a DOCX, in memory. The document itself is never modified."""
from __future__ import annotations

import io
import re


class DocxTextError(RuntimeError):
    pass


def _import_docx():
    try:
        import docx  # noqa: PLC0415
    except ImportError as exc:  # pragma: no cover - dependency guard
        raise SystemExit("ERROR: python-docx is required. Run: pip install -r requirements.txt") from exc
    return docx


def normalize_text(text: str, *, keep_columns: bool = False) -> str:
    text = (text or "").replace("\r\n", "\n").replace("\r", "\n")
    text = text.replace("\u00a0", " ").replace("\u2013", "-").replace("\u2014", "-")
    text = re.sub(r"\t", " ", text)
    text = re.sub(r" {2,}", "  " if keep_columns else " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return "\n".join(line.strip() for line in text.split("\n")).strip()


def looks_empty(text: str, *, min_chars: int = 40) -> bool:
    return len(re.sub(r"\s", "", text or "")) < min_chars


def extract_text(docx_bytes: bytes) -> str:
    """Return text from paragraphs and tables in reading order as best effort."""
    docx = _import_docx()
    try:
        document = docx.Document(io.BytesIO(docx_bytes))
    except Exception as exc:
        raise DocxTextError(f"Could not open DOCX: {exc}") from exc

    chunks: list[str] = []
    for paragraph in document.paragraphs:
        text = (paragraph.text or "").strip()
        if text:
            chunks.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [(cell.text or "").strip() for cell in row.cells]
            cells = [cell for cell in cells if cell]
            if cells:
                chunks.append(" | ".join(cells))

    return normalize_text("\n".join(chunks))


def section_count(docx_bytes: bytes) -> int:
    docx = _import_docx()
    try:
        document = docx.Document(io.BytesIO(docx_bytes))
    except Exception:
        return 0
    return len(document.sections)
