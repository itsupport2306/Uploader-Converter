"""
screenshot_to_word.py
=====================
Convert Doximity-style physician-profile screenshots or NPPES-style Excel rows
into clean, editable Word (.docx) documents.

The tool:
  1. De-inverts the dark header band so white-on-dark text (the name) is read.
  2. Detects the column gap and drops the right sidebar (contact info + the
     "Similar Physicians & HCPs" list) plus join/promo call-to-actions.
  3. Rebuilds document structure (name, specialty, sections, and the entries
     inside each section) from OCR word boxes.
  4. Renders a formatted .docx that mirrors the on-screen layout.

Usage:
    python screenshot_to_word.py INPUT [-o OUTPUT] [options]

INPUT may be a single image, a folder, or an Excel/CSV file. Image inputs use
OCR; tabular inputs write one Word document per data row.
Run with --debug to print the detected structure for tuning.
"""

from __future__ import annotations

import argparse
import hashlib
import csv
import json
import os
import re
import sys
import tempfile
import time
import webbrowser
from dataclasses import dataclass, field
from datetime import date, datetime
from pathlib import Path
from urllib.error import HTTPError, URLError
from urllib.parse import quote, urlencode
from urllib.request import Request, urlopen
from zipfile import BadZipFile

import numpy as np
import pytesseract
from PIL import Image, ImageOps
from pytesseract import Output

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

# --------------------------------------------------------------------------- #
# Configuration
# --------------------------------------------------------------------------- #

# Tesseract binary locations to probe (in order) if not already on PATH.
TESSERACT_CANDIDATES = [
    r"C:\Program Files\Tesseract-OCR\tesseract.exe",
    r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
    os.path.expandvars(r"%LOCALAPPDATA%\Programs\Tesseract-OCR\tesseract.exe"),
]

IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tif", ".tiff"}
EXCEL_EXTS = {".xlsx", ".xlsm"}
CSV_EXTS = {".csv"}
TABULAR_EXTS = EXCEL_EXTS | CSV_EXTS
GRAPH_BASE_URL = "https://graph.microsoft.com/v1.0"
MICROSOFT_LOGIN_BASE_URL = "https://login.microsoftonline.com"
ONEDRIVE_SCOPES = "Files.ReadWrite offline_access User.Read"

NPPES_PART01_HEADERS = [
    "NPI",
    "Prefix",
    "First Name",
    "Middle Name",
    "Last Name",
    "Suffix",
    "Gender",
    "Credential",
    "Birth Date",
    "Provider Type",
    "Specialty",
    "Taxonomy Code",
    "Address Type",
    "Street Address",
    "Street Address 2",
    "City",
    "State",
    "Zip Code",
    "Country",
    "Contact Type",
    "Phone Number",
    "School Name",
    "Graduation Year",
    "License Number",
    "License State",
    "Enumeration Date",
    "Last Update",
]

# Canonical section headings seen on Doximity profiles. Matching is fuzzy
# (lower-cased, punctuation-insensitive, "contains") so OCR noise is tolerated.
# NOTE: only distinctive multi-word headings are listed. Generic single words
# like "Experience" / "Specialties" are intentionally excluded — they collide
# with wrapped publication titles ("...Single Center Real-World Experience").
KNOWN_HEADINGS = [
    "Education & Training",
    "Certifications & Licensure",
    "Awards, Honors, & Recognition",
    "Publications & Presentations",
    "Professional Memberships",
    "Languages",
]

# Sub-headings rendered smaller/bold inside a section (e.g. the source of a
# publication list). Kept as bold lead-ins rather than top-level headings.
KNOWN_SUBHEADINGS = ["PubMed", "Other"]

# Promotional / navigational lines that belong to Doximity's chrome, not the
# physician's data. Dropped unless --keep-promo is given. Matched as a
# normalized "contains" against the line text.
PROMO_SNIPPETS = [
    "oximity",               # the brand logo, incl. OCR variants ("Aoximity")
    "join to view full profile",
    "is on doximity",
    "as a doximity member",
    "gain access to free telehealth",
    "connect with colleagues",
    "read the latest clinical news",
    "full profile",          # "See Dr. X's Full Profile"
    "already have an account",
    "viewing the full profile is available",
    "verified healthcare professionals only",
    "similar physicians",
    "join for free",
]

# Markers of Doximity's page footer / SEO block, which appears below the real
# profile content on shorter pages. Once one is seen, everything after it is
# dropped (it's site chrome: nav links, app-store text, name-variant SEO).
FOOTER_SNIPPETS = [
    "about us",
    "investors",
    "trademarks of apple",
    "iphone and ipad",
    "download on the app store",
    "google play",
    "terms of service",
    "privacy policy",
]

# Base typography for the generated document.
BASE_FONT = "Calibri"
COLOR_TEXT = RGBColor(0x20, 0x20, 0x20)      # near-black body
COLOR_MUTED = RGBColor(0x5F, 0x63, 0x68)     # gray detail lines
COLOR_HEADING = RGBColor(0x14, 0x14, 0x14)
COLOR_ACCENT = RGBColor(0x0B, 0x57, 0xD0)     # link/specialty blue


# --------------------------------------------------------------------------- #
# Data model
# --------------------------------------------------------------------------- #

@dataclass
class Line:
    text: str
    left: int
    top: int
    right: int
    bottom: int
    height: int          # median glyph height (font-size proxy)

    @property
    def width(self) -> int:
        return self.right - self.left


@dataclass
class Block:
    """A rendered unit: a kind plus its text and (for entries) detail lines."""
    kind: str            # title | subtitle | heading | subheading | entry | body
    text: str
    details: list[str] = field(default_factory=list)


@dataclass
class ConversionSummary:
    converted: int = 0
    skipped: int = 0
    failed: int = 0
    total: int = 0


# --------------------------------------------------------------------------- #
# Tesseract setup
# --------------------------------------------------------------------------- #

def configure_tesseract(explicit: str | None = None) -> None:
    """Point pytesseract at a working tesseract binary or exit with guidance."""
    if explicit:
        pytesseract.pytesseract.tesseract_cmd = explicit
    else:
        from shutil import which
        found = which("tesseract")
        if not found:
            for cand in TESSERACT_CANDIDATES:
                if os.path.isfile(cand):
                    found = cand
                    break
        if found:
            pytesseract.pytesseract.tesseract_cmd = found

    try:
        pytesseract.get_tesseract_version()
    except Exception:  # pragma: no cover - environment dependent
        sys.exit(
            "ERROR: Tesseract OCR was not found.\n"
            "Install it (Windows build: https://github.com/UB-Mannheim/tesseract/wiki)\n"
            "then re-run, optionally passing --tesseract \"C:\\path\\to\\tesseract.exe\"."
        )


# --------------------------------------------------------------------------- #
# Image preprocessing
# --------------------------------------------------------------------------- #

def load_image(path: Path, scale: float) -> Image.Image:
    img = Image.open(path).convert("RGB")
    if scale and abs(scale - 1.0) > 1e-3:
        img = img.resize(
            (int(img.width * scale), int(img.height * scale)), Image.LANCZOS
        )
    return img


def deinvert_dark_header(img: Image.Image, dark_thresh: int = 110) -> tuple[Image.Image, int]:
    """Invert a contiguous dark band at the top so light-on-dark text is read.

    Returns the (possibly modified) image and the y where the header ends.
    """
    gray = np.asarray(img.convert("L"))
    row_median = np.median(gray, axis=1)
    h = gray.shape[0]

    end = -1
    gap = 0
    for y in range(h):
        if row_median[y] < dark_thresh:
            end = y
            gap = 0
        else:
            gap += 1
            # Allow small light gaps inside the header before deciding it ended.
            if end >= 0 and gap > max(15, h // 100):
                break
        # Header should not span more than ~30% of the image.
        if y > h * 0.30:
            break

    if end <= 0:
        return img, 0

    out = img.copy()
    band = out.crop((0, 0, out.width, end + 1))
    out.paste(ImageOps.invert(band), (0, 0))
    return out, end + 1


# --------------------------------------------------------------------------- #
# OCR + layout analysis
# --------------------------------------------------------------------------- #

def ocr_words(img: Image.Image, psm: int = 3, min_conf: int = 35) -> list[dict]:
    data = pytesseract.image_to_data(
        img, output_type=Output.DICT, config=f"--psm {psm}"
    )
    words = []
    for i, text in enumerate(data["text"]):
        if not text.strip():
            continue
        try:
            conf = float(data["conf"][i])
        except (ValueError, TypeError):
            conf = -1
        if conf < min_conf:
            continue
        words.append(
            {
                "text": text,
                "left": data["left"][i],
                "top": data["top"][i],
                "width": data["width"][i],
                "height": data["height"][i],
                "block": data["block_num"][i],
                "par": data["par_num"][i],
                "line": data["line_num"][i],
            }
        )
    return words


def find_column_cutoff(
    words: list[dict], img_w: int, header_end: int, default_ratio: float
) -> int:
    """Find the x splitting the main column from the right sidebar.

    Looks for the widest vertical whitespace gap in the right half of the body
    region. Falls back to ``default_ratio`` * width when detection is unclear.
    """
    fallback = int(img_w * default_ratio)
    coverage = np.zeros(img_w + 1, dtype=np.int32)
    for w in words:
        if w["top"] < header_end:          # ignore the header band
            continue
        l = max(0, w["left"])
        r = min(img_w, w["left"] + w["width"])
        if r > l:
            coverage[l:r] += 1

    if coverage.sum() == 0:
        return fallback

    covered = coverage > 0
    search_start = int(img_w * 0.45)        # sidebars live in the right portion
    min_gap = max(22, img_w // 60)

    run_start = None
    for x in range(search_start, img_w):
        if not covered[x]:
            if run_start is None:
                run_start = x
        else:
            if run_start is not None and x - run_start >= min_gap:
                # First wide gap after the dense main column = the divider.
                if coverage[:run_start].sum() > coverage[run_start:].sum() * 0.4:
                    return run_start
            run_start = None
    return fallback


def clean_text(s: str) -> str:
    """Normalize OCR text: map mis-read bullet glyphs and tidy whitespace."""
    # Drop the "N citations" badge (and its § icon) that sits beside pub titles.
    s = re.sub(r"\s*§?\s*\d{1,5}\s+citations?\b", "", s, flags=re.I)
    s = s.replace("§", "")
    # A "•" separator is frequently mis-read as «, », ·, or the replacement char.
    s = re.sub(r"\s*[«»‹›·�]\s*", " • ", s)
    s = s.replace("|", "")                        # stray vertical rules
    # A symbol glued to the front of the first real word is an icon artifact
    # (e.g. "{Albert" -> "Albert"); real lines never start with these.
    s = re.sub(r"^[\[\]{}<>~^`\\=]+\s*(?=[A-Za-z0-9])", "", s)
    s = re.sub(r"\s+", " ", s).strip(" •")
    return s.strip()


def _is_punct(token: str) -> bool:
    return re.sub(r"[\W_]", "", token) == ""


def group_lines(words: list[dict], header_end: int = 0) -> list[Line]:
    """Group OCR words into lines, dropping leading left-margin icon artifacts.

    Each Doximity entry has a small logo/icon in the left margin that OCR turns
    into 1-4 character noise (e.g. "Hy", "lg)", "*"). Such tokens sit left of the
    text column, so we detect the body text indent and strip leading short tokens
    that fall in the icon strip (or are pure punctuation).
    """
    buckets: dict[tuple[int, int, int], list[dict]] = {}
    for w in words:
        buckets.setdefault((w["block"], w["par"], w["line"]), []).append(w)

    raw = []  # (sorted_words, joined_text)
    for key in sorted(buckets):
        ws = sorted(buckets[key], key=lambda w: w["left"])
        txt = clean_text(" ".join(w["text"] for w in ws))
        if txt:
            raw.append((ws, txt))

    if not raw:
        return []

    body_h = int(np.median([np.median([w["height"] for w in ws]) for ws, _ in raw]))

    lines: list[Line] = []
    for ws, _ in raw:
        ws = list(ws)
        # Header lines use large fonts; don't mistake big title words for logos.
        in_header = min(w["top"] for w in ws) < header_end
        while ws and not in_header:
            w0 = ws[0]
            short = len(re.sub(r"\W", "", w0["text"])) <= 4
            if _is_punct(w0["text"]):
                ws = ws[1:]
                continue
            # A short token carrying OCR-only symbols (e.g. "R==~") is icon noise;
            # these symbols never appear in real names, codes, or dates.
            if short and any(c in "=~{}<>|^\\`" for c in w0["text"]):
                ws = ws[1:]
                continue
            if not short:
                break
            if len(ws) == 1:
                # A lone short token whose glyph is far from body size is a logo
                # (tiny rendered mark like "ABMS", or an oversized icon).
                h = w0["height"]
                if h < 0.55 * body_h or h > 2.2 * body_h:
                    ws = ws[1:]
                break
            # An entry's left-margin logo OCRs as a short token that is either
            # much taller than body text (a graphic) or far from the next word.
            gap = ws[1]["left"] - (w0["left"] + w0["width"])
            tall = w0["height"] > 1.5 * body_h
            big_gap = gap > 1.2 * body_h
            if tall or big_gap:
                ws = ws[1:]
            else:
                break
        if not ws:
            continue
        text = clean_text(" ".join(w["text"] for w in ws))
        if not text:
            continue
        left = min(w["left"] for w in ws)
        top = min(w["top"] for w in ws)
        right = max(w["left"] + w["width"] for w in ws)
        bottom = max(w["top"] + w["height"] for w in ws)
        height = int(np.median([w["height"] for w in ws]))
        lines.append(Line(text, left, top, right, bottom, height))

    lines.sort(key=lambda ln: (ln.top, ln.left))
    return lines


# --------------------------------------------------------------------------- #
# Structure reconstruction
# --------------------------------------------------------------------------- #

def _norm(s: str) -> str:
    return re.sub(r"[^a-z0-9 ]", "", s.lower()).strip()


def is_promo(text: str) -> bool:
    n = _norm(text)
    return any(snippet in n for snippet in (_norm(p) for p in PROMO_SNIPPETS))


def is_footer(text: str) -> bool:
    """True for Doximity footer / SEO lines that mark the end of real content."""
    n = _norm(text)
    if any(snippet in n for snippet in (_norm(p) for p in FOOTER_SNIPPETS)):
        return True
    # SEO name-variant line, e.g. "Dr. X, Dr. X MD, Dr. E X, Dr. X Michael".
    return text.count("Dr.") >= 3


_NAME_CRED = re.compile(r"\b(MD|DO|MBBS|MBChB|DDS|DMD|DPM|DPT|DC|ND|PharmD)\b")


def _looks_like_name(text: str) -> bool:
    """A profile name line carries a degree credential near its end (e.g. '... MD')."""
    matches = list(_NAME_CRED.finditer(text))
    return bool(matches) and matches[-1].start() >= len(text) * 0.45


def _strip_leading_junk(text: str) -> str:
    """Remove a leading icon/OCR artifact from a header line (punctuation run or
    a short lowercase fragment sitting before the first real, capitalized word)."""
    text = re.sub(r"^\s*[^\w\s]+\s*", "", text)
    text = re.sub(r"^([a-z]{1,3})\s+(?=[A-Z0-9])", "", text)
    return text.strip()


def match_heading(text: str) -> str | None:
    n = _norm(text)
    if len(n) < 3:
        return None
    for h in KNOWN_HEADINGS:
        hn = _norm(h)
        if n == hn or (len(n) <= len(hn) + 6 and hn in n) or (n in hn and len(n) > len(hn) - 4):
            return h
    return None


def match_subheading(text: str) -> str | None:
    n = _norm(text)
    for h in KNOWN_SUBHEADINGS:
        if n == _norm(h):
            return h
    return None


def build_blocks(lines: list[Line], header_end: int, keep_promo: bool) -> list[Block]:
    """Turn classified lines into ordered rendering blocks."""
    if not lines:
        return []

    body = [ln for ln in lines if ln.bottom > header_end]
    body_height = int(np.median([ln.height for ln in body])) if body else 12

    blocks: list[Block] = []

    # --- Header: title + subtitle lines (everything above the body region) ---
    header_lines = [ln for ln in lines if ln.top < header_end]
    if not keep_promo:
        header_lines = [ln for ln in header_lines if not is_promo(ln.text)]
    header_lines = [ln for ln in header_lines if len(_norm(ln.text)) >= 2]
    if header_lines:
        # The name is the header line that carries a degree credential near its
        # end ("... MD"); this beats "tallest", which can pick a stray logo glyph
        # (e.g. "ov", "“Aoximity") or a "MD at <employer>" line. Fall back to the
        # tallest line only if no credentialed name line is present.
        name_lines = [ln for ln in header_lines if _looks_like_name(ln.text)]
        if name_lines:
            title_line = max(name_lines, key=lambda ln: (ln.height, len(ln.text)))
        else:
            title_line = max(header_lines, key=lambda ln: ln.height)
        blocks.append(Block("title", _strip_leading_junk(title_line.text)))
        for ln in header_lines:
            if ln is title_line:
                continue
            sub = _strip_leading_junk(ln.text)
            if len(_norm(sub)) > 3:        # skip tiny logo/icon remnants ("ov")
                blocks.append(Block("subtitle", sub))

    # --- Body: sections, sub-headings, and gap-separated entries -------------
    seen_heading = False
    pending: list[Line] = []          # lines accumulating into the current entry group

    def flush_entries(group: list[Line]):
        """Split a run of body lines into entries by vertical spacing."""
        if not group:
            return
        gaps = [
            group[i].top - group[i - 1].bottom for i in range(1, len(group))
        ]
        # An "entry break" is a gap noticeably larger than the typical line gap.
        typical = np.median(gaps) if gaps else 0
        threshold = max(typical * 1.8, body_height * 0.9)

        entry: list[Line] = [group[0]]
        for i in range(1, len(group)):
            gap = group[i].top - group[i - 1].bottom
            if gap > threshold:
                _emit_entry(entry)
                entry = [group[i]]
            else:
                entry.append(group[i])
        _emit_entry(entry)

    def _emit_entry(entry: list[Line]):
        if not entry:
            return
        head = entry[0].text
        details = [ln.text for ln in entry[1:]]
        blocks.append(Block("entry", head, details))

    for ln in body:
        if not keep_promo and is_footer(ln.text):
            break  # site footer / SEO block — nothing real comes after it
        if not keep_promo and is_promo(ln.text):
            continue

        heading = match_heading(ln.text)
        if heading and ln.height >= body_height * 0.9:
            flush_entries(pending)
            pending = []
            blocks.append(Block("heading", heading))
            seen_heading = True
            continue

        sub = match_subheading(ln.text)
        if sub and seen_heading:
            flush_entries(pending)
            pending = []
            blocks.append(Block("subheading", sub))
            continue

        if not seen_heading:
            # Everything before the first section heading is Doximity's join /
            # promo box (the physician's data starts at the first heading).
            if keep_promo:
                blocks.append(Block("body", ln.text))
            continue

        pending.append(ln)

    flush_entries(pending)
    return blocks


# --------------------------------------------------------------------------- #
# DOCX rendering
# --------------------------------------------------------------------------- #

def _set_run(run, *, size=None, bold=False, italic=False, color=None, font=BASE_FONT):
    run.font.name = font
    if size is not None:
        run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def _add_bottom_border(paragraph):
    """Draw a thin rule under a paragraph to mimic Doximity's section dividers."""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "D0D3D8")
    borders.append(bottom)
    pPr.append(borders)


def render_docx(blocks: list[Block], out_path: Path) -> None:
    doc = Document()

    # Base style.
    normal = doc.styles["Normal"]
    normal.font.name = BASE_FONT
    normal.font.size = Pt(11)
    normal.font.color.rgb = COLOR_TEXT

    for block in blocks:
        if block.kind == "title":
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            _set_run(p.add_run(block.text), size=22, bold=True, color=COLOR_HEADING)

        elif block.kind == "subtitle":
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            _set_run(p.add_run(block.text), size=12, color=COLOR_ACCENT)

        elif block.kind == "heading":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(6)
            _set_run(p.add_run(block.text), size=15, bold=True, color=COLOR_HEADING)
            _add_bottom_border(p)

        elif block.kind == "subheading":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            _set_run(p.add_run(block.text), size=12, bold=True, color=COLOR_HEADING)

        elif block.kind == "entry":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(0)
            _set_run(p.add_run(block.text), size=11.5, bold=True, color=COLOR_TEXT)
            for detail in block.details:
                d = doc.add_paragraph()
                d.paragraph_format.space_after = Pt(0)
                _set_run(d.add_run(detail), size=10.5, color=COLOR_MUTED)

        elif block.kind == "body":
            p = doc.add_paragraph()
            _set_run(p.add_run(block.text), size=11, color=COLOR_TEXT)

    doc.save(str(out_path))


# --------------------------------------------------------------------------- #
# Tabular / Excel reconstruction
# --------------------------------------------------------------------------- #

def _norm_key(s: str) -> str:
    return re.sub(r"[^a-z0-9]+", "", s.lower())


def _cell_to_text(value) -> str:
    if value is None:
        return ""
    if isinstance(value, bool):
        return "Yes" if value else "No"
    if isinstance(value, int):
        return str(value)
    if isinstance(value, float):
        return str(int(value)) if value.is_integer() else f"{value:g}"
    if isinstance(value, (datetime, date)):
        return value.isoformat()
    return re.sub(r"\s+", " ", str(value)).strip()


def _unique_headers(raw_headers, width: int) -> list[str]:
    headers = [_cell_to_text(h) for h in raw_headers]
    if width == len(NPPES_PART01_HEADERS):
        headers = [
            text or NPPES_PART01_HEADERS[i]
            for i, text in enumerate(headers[:width])
        ]

    out: list[str] = []
    seen: dict[str, int] = {}
    for i in range(width):
        text = headers[i] if i < len(headers) else ""
        if not text:
            text = f"Column {i + 1}"
        key = _norm_key(text) or f"column{i + 1}"
        seen[key] = seen.get(key, 0) + 1
        out.append(text if seen[key] == 1 else f"{text} {seen[key]}")
    return out


def _lookup(row: dict[str, str]) -> dict[str, str]:
    return {_norm_key(k): v for k, v in row.items() if v}


def _get(lookup: dict[str, str], *names: str) -> str:
    for name in names:
        value = lookup.get(_norm_key(name))
        if value:
            return value
    return ""


def _join_parts(parts) -> str:
    return " ".join(part for part in parts if part)


def _title_case_if_all_caps(text: str) -> str:
    if text and text == text.upper() and any(ch.isalpha() for ch in text):
        return text.title()
    return text


def _entity_type_label(code: str) -> str:
    code = code.strip()
    if code == "1":
        return "Individual"
    if code == "2":
        return "Organization"
    return code


def _format_phone(value: str) -> str:
    digits = re.sub(r"\D", "", value)
    if len(digits) == 10:
        return f"({digits[:3]}) {digits[3:6]}-{digits[6:]}"
    return value


def _format_postal(value: str) -> str:
    digits = re.sub(r"\D", "", value)
    if len(digits) == 9:
        return f"{digits[:5]}-{digits[5:]}"
    if digits and len(digits) < 5:
        return digits.zfill(5)
    return value


def _city_state_zip(city: str, state: str, postal: str) -> str:
    left = ", ".join(part for part in [city, state] if part)
    if postal:
        return f"{left} {_format_postal(postal)}".strip()
    return left


def _provider_display_name(lookup: dict[str, str]) -> str:
    org = _get(
        lookup,
        "Provider Organization Name (Legal Business Name)",
        "Provider Organization Name",
        "Organization Name",
    )
    if org:
        return _title_case_if_all_caps(org)

    name = _join_parts(
        [
            _title_case_if_all_caps(_get(lookup, "Provider Name Prefix Text", "Name Prefix", "Prefix")),
            _title_case_if_all_caps(_get(lookup, "Provider First Name", "First Name")),
            _title_case_if_all_caps(_get(lookup, "Provider Middle Name", "Middle Name")),
            _title_case_if_all_caps(
                _get(
                    lookup,
                    "Provider Last Name (Legal Name)",
                    "Provider Last Name",
                    "Last Name",
                )
            ),
            _title_case_if_all_caps(_get(lookup, "Provider Name Suffix Text", "Name Suffix", "Suffix")),
        ]
    )
    name = _title_case_if_all_caps(name)
    credential = _get(lookup, "Provider Credential Text", "Credential")
    if name and credential and credential.casefold() not in name.casefold().split():
        name = f"{name} {credential}"
    return name


def _other_provider_name(lookup: dict[str, str]) -> tuple[str, list[str]]:
    other_org = _get(lookup, "Provider Other Organization Name", "Other Organization Name")
    if other_org:
        details = []
        name_type = _get(lookup, "Provider Other Organization Name Type Code")
        if name_type:
            details.append(f"Type Code: {name_type}")
        return other_org, details

    other_name = _join_parts(
        [
            _get(lookup, "Provider Other Name Prefix Text"),
            _get(lookup, "Provider Other First Name"),
            _get(lookup, "Provider Other Middle Name"),
            _get(lookup, "Provider Other Last Name"),
            _get(lookup, "Provider Other Name Suffix Text"),
        ]
    )
    details = []
    credential = _get(lookup, "Provider Other Credential Text")
    name_type = _get(lookup, "Provider Other Last Name Type Code")
    if credential:
        details.append(f"Credential: {credential}")
    if name_type:
        details.append(f"Type Code: {name_type}")
    return other_name, details


def _is_nppes_row(lookup: dict[str, str]) -> bool:
    return bool(
        _get(lookup, "NPI")
        or _get(lookup, "Provider First Name", "First Name")
        or _get(lookup, "Provider Last Name (Legal Name)", "Last Name")
        or _get(lookup, "Provider Organization Name (Legal Business Name)")
    )


def tabular_row_to_blocks(
    row: dict[str, str], *, row_number: int, source_name: str
) -> list[Block]:
    lookup = _lookup(row)
    if not _is_nppes_row(lookup):
        title = _get(lookup, "Name", "Full Name", "Provider Name") or f"Excel Row {row_number}"
        blocks = [Block("title", title), Block("subtitle", source_name)]
        blocks.append(Block("heading", "Excel Information"))
        for key, value in row.items():
            if value:
                blocks.append(Block("entry", key, [value]))
        return blocks

    npi = _get(lookup, "NPI")
    entity_type = _entity_type_label(_get(lookup, "Entity Type Code", "Entity Type"))
    provider_type = _get(lookup, "Provider Type")
    specialty = _get(lookup, "Specialty")
    taxonomy = _get(lookup, "Taxonomy Code")
    city = _title_case_if_all_caps(_get(lookup, "Provider Business Mailing Address City Name", "City"))
    state = _get(lookup, "Provider Business Mailing Address State Name", "State")
    postal = _get(lookup, "Provider Business Mailing Address Postal Code", "Postal Code", "Zip Code")
    location = _city_state_zip(city, state, postal)

    title = _provider_display_name(lookup) or (f"NPI {npi}" if npi else f"Excel Row {row_number}")
    subtitle_parts = [
        part for part in [specialty or provider_type or entity_type, location] if part
    ]

    blocks = [Block("title", title)]
    if subtitle_parts:
        blocks.append(Block("subtitle", " * ".join(subtitle_parts)))

    info_details = []
    if npi:
        info_details.append(f"NPI: {npi}")
    if entity_type:
        info_details.append(f"Entity Type: {entity_type}")
    if provider_type:
        info_details.append(f"Provider Type: {provider_type}")
    if specialty:
        info_details.append(f"Specialty: {specialty}")
    if taxonomy:
        info_details.append(f"Taxonomy Code: {taxonomy}")
    gender = _get(lookup, "Gender")
    if gender:
        info_details.append(f"Gender: {gender}")
    birth_date = _get(lookup, "Birth Date")
    if birth_date:
        info_details.append(f"Birth Date: {birth_date}")
    credential = _get(lookup, "Provider Credential Text", "Credential")
    if credential:
        info_details.append(f"Credential: {credential}")
    blocks.append(Block("heading", "Provider Information"))
    blocks.append(Block("entry", title, info_details or ["NPPES provider record"]))

    other_name, other_details = _other_provider_name(lookup)
    if other_name:
        blocks.append(Block("heading", "Other Names"))
        blocks.append(Block("entry", other_name, other_details))

    address_type = _get(lookup, "Address Type")
    address1 = _get(
        lookup,
        "Provider First Line Business Mailing Address",
        "Address Line 1",
        "Street Address",
    )
    address2 = _get(
        lookup,
        "Provider Second Line Business Mailing Address",
        "Address Line 2",
        "Street Address 2",
    )
    country = _get(
        lookup,
        "Provider Business Mailing Address Country Code (If outside U.S.)",
        "Country",
    )
    address_details = []
    if address2:
        address_details.append(_title_case_if_all_caps(address2))
    if location:
        address_details.append(location)
    if country and _norm_key(country) not in {"us", "usa", "unitedstates"}:
        address_details.append(country)
    if address1 or address_details:
        heading = f"{address_type} Address" if address_type else "Business Mailing Address"
        blocks.append(Block("heading", heading))
        blocks.append(Block("entry", _title_case_if_all_caps(address1) or heading, address_details))

    school_name = _get(lookup, "School Name")
    graduation_year = _get(lookup, "Graduation Year")
    if school_name or graduation_year:
        blocks.append(Block("heading", "Education & Training"))
        blocks.append(
            Block(
                "entry",
                _title_case_if_all_caps(school_name) or "Education",
                [f"Graduation Year: {graduation_year}"] if graduation_year else [],
            )
        )

    license_number = _get(lookup, "License Number")
    license_state = _get(lookup, "License State")
    if license_number or license_state:
        blocks.append(Block("heading", "Certifications & Licensure"))
        title_text = f"{license_state} State License" if license_state else "State License"
        license_details = []
        if license_number:
            license_details.append(f"License Number: {license_number}")
        blocks.append(Block("entry", title_text, license_details))

    identifier_details = []
    replacement_npi = _get(lookup, "Replacement NPI")
    ein = _get(lookup, "Employer Identification Number (EIN)", "EIN")
    if replacement_npi:
        identifier_details.append(f"Replacement NPI: {replacement_npi}")
    if ein:
        identifier_details.append(f"EIN: {ein}")
    if identifier_details:
        blocks.append(Block("heading", "Identifiers"))
        blocks.append(Block("entry", "Additional Identifiers", identifier_details))

    enumeration_date = _get(lookup, "Enumeration Date")
    last_update = _get(lookup, "Last Update")
    if enumeration_date or last_update:
        registry_details = []
        if enumeration_date:
            registry_details.append(f"Enumeration Date: {enumeration_date}")
        if last_update:
            registry_details.append(f"Last Update: {last_update}")
        blocks.append(Block("heading", "Registry Dates"))
        blocks.append(Block("entry", "NPPES Record", registry_details))

    return blocks


def iter_csv_rows(input_path: Path):
    with input_path.open("r", encoding="utf-8-sig", newline="") as fh:
        reader = csv.DictReader(fh)
        for row_number, row in enumerate(reader, start=2):
            clean = {key or "Column": _cell_to_text(value) for key, value in row.items()}
            if any(clean.values()):
                yield row_number, clean


def iter_workbook_rows(input_path: Path, sheet_name: str | None = None):
    try:
        from openpyxl import load_workbook
    except ImportError as exc:  # pragma: no cover - depends on environment setup
        raise RuntimeError(
            "Excel conversion requires openpyxl. Install dependencies with "
            "`pip install -r requirements.txt`."
        ) from exc

    try:
        wb = load_workbook(input_path, read_only=True, data_only=True)
    except (BadZipFile, EOFError) as exc:
        raise RuntimeError(
            f"{input_path} is not a readable .xlsx/.xlsm workbook. The file may "
            "be incomplete or corrupted; re-export or re-download it and try again."
        ) from exc

    try:
        if sheet_name:
            if sheet_name not in wb.sheetnames:
                raise RuntimeError(
                    f"Worksheet {sheet_name!r} was not found. Available sheets: "
                    f"{', '.join(wb.sheetnames)}"
                )
            ws = wb[sheet_name]
        else:
            ws = wb.active

        rows = ws.iter_rows(values_only=True)
        header_values = next(rows, None)
        if header_values is None:
            return
        width = len(header_values)
        headers = _unique_headers(header_values, width)

        for row_number, values in enumerate(rows, start=2):
            values = values[:width] if len(values) > width else values
            clean_values = [_cell_to_text(value) for value in values]
            if not any(clean_values):
                continue
            row = {
                headers[i]: clean_values[i] if i < len(clean_values) else ""
                for i in range(width)
            }
            yield row_number, row
    finally:
        wb.close()


def iter_tabular_rows(input_path: Path, sheet_name: str | None = None):
    if input_path.suffix.lower() in CSV_EXTS:
        yield from iter_csv_rows(input_path)
    else:
        yield from iter_workbook_rows(input_path, sheet_name)


def _safe_filename_part(text: str, max_len: int = 80) -> str:
    text = re.sub(r'[<>:"/\\|?*\x00-\x1F]+', "_", text)
    text = re.sub(r"\s+", "_", text).strip("._ ")
    if len(text) > max_len:
        text = text[:max_len].rstrip("._ ")
    return text


def tabular_output_path(row: dict[str, str], row_number: int, out_dir: Path) -> Path:
    lookup = _lookup(row)
    name = _provider_display_name(lookup) or _get(lookup, "Name", "Full Name")
    npi = _get(lookup, "NPI")
    if name:
        stem = _safe_filename_part(name)
        if npi:
            stem = f"{stem}_{_safe_filename_part(npi, max_len=20)}"
        else:
            stem = f"{stem}_row_{row_number}"
    else:
        stem = f"row_{row_number}"
    return out_dir / f"{stem or f'row_{row_number}'}.docx"


def resume_state_path(
    input_path: Path,
    *,
    output_identity: str,
    sheet_name: str | None,
) -> Path:
    try:
        resolved = str(input_path.resolve())
    except OSError:
        resolved = str(input_path)
    key = json.dumps(
        {
            "input": resolved.lower(),
            "output": output_identity,
            "sheet": sheet_name or "",
        },
        sort_keys=True,
    )
    digest = hashlib.sha1(key.encode("utf-8")).hexdigest()[:12]
    state_dir = Path.cwd() / ".converter_state"
    return state_dir / f"{input_path.stem}_{digest}.json"


def load_resume_row(state_path: Path, input_path: Path) -> int | None:
    if not state_path.exists():
        return None
    try:
        state = json.loads(state_path.read_text(encoding="utf-8"))
        current_mtime = input_path.stat().st_mtime
        if abs(float(state.get("input_mtime", -1)) - current_mtime) > 1:
            return None
        next_row = int(state.get("next_row", 0))
        return next_row if next_row >= 2 else None
    except Exception:
        return None


def save_resume_row(state_path: Path, input_path: Path, next_row: int) -> None:
    state_path.parent.mkdir(parents=True, exist_ok=True)
    state = {
        "input": str(input_path),
        "input_mtime": input_path.stat().st_mtime,
        "next_row": next_row,
        "updated_at": datetime.now().isoformat(timespec="seconds"),
    }
    tmp_path = state_path.with_suffix(state_path.suffix + ".tmp")
    tmp_path.write_text(json.dumps(state, indent=2), encoding="utf-8")
    os.replace(tmp_path, state_path)


def _write_docx_atomic(blocks: list[Block], out_path: Path) -> None:
    tmp_path = out_path.parent / (out_path.name + ".tmp")
    render_docx(blocks, tmp_path)
    os.replace(tmp_path, out_path)


class OneDriveError(RuntimeError):
    def __init__(
        self,
        message: str,
        *,
        status: int | None = None,
        payload=None,
        retry_after: int | None = None,
    ):
        super().__init__(message)
        self.status = status
        self.payload = payload or {}
        self.retry_after = retry_after


def _read_http_error(exc: HTTPError) -> dict:
    try:
        raw = exc.read().decode("utf-8", errors="replace")
    except Exception:
        raw = ""
    try:
        payload = json.loads(raw) if raw else {}
    except json.JSONDecodeError:
        payload = {"error_description": raw}
    if isinstance(payload.get("error"), dict):
        message = payload["error"].get("message") or str(payload["error"])
    else:
        message = payload.get("error_description") or payload.get("error") or raw
    return {"payload": payload, "message": message or str(exc)}


def _request_json(req: Request, *, ok_statuses: set[int] | None = None) -> dict:
    ok_statuses = ok_statuses or {200, 201}
    try:
        with urlopen(req, timeout=60) as resp:
            body = resp.read().decode("utf-8", errors="replace")
            if resp.status not in ok_statuses:
                raise OneDriveError(f"HTTP {resp.status}: {body}", status=resp.status)
            return json.loads(body) if body else {}
    except HTTPError as exc:
        parsed = _read_http_error(exc)
        retry_after = None
        header_value = exc.headers.get("Retry-After") if exc.headers else None
        if header_value and header_value.isdigit():
            retry_after = int(header_value)
        raise OneDriveError(
            parsed["message"],
            status=exc.code,
            payload=parsed["payload"],
            retry_after=retry_after,
        ) from exc
    except URLError as exc:
        raise OneDriveError(f"Could not reach Microsoft endpoint: {exc}") from exc


def _post_form(url: str, fields: dict[str, str]) -> dict:
    data = urlencode(fields).encode("utf-8")
    req = Request(
        url,
        data=data,
        headers={"Content-Type": "application/x-www-form-urlencoded"},
        method="POST",
    )
    return _request_json(req)


class OneDriveUploader:
    def __init__(
        self,
        *,
        client_id: str,
        tenant: str,
        remote_folder: str,
        upload_delay: float = 0.0,
        open_browser: bool = True,
    ):
        self.client_id = client_id
        self.tenant = tenant.strip().strip("/") or "consumers"
        self.remote_folder = self._normalize_remote_folder(remote_folder)
        self.upload_delay = max(0.0, upload_delay)
        self.open_browser = open_browser
        self.access_token = ""
        self.refresh_token = ""
        self.expires_at = 0.0
        self._folder_ready = False

    @property
    def oauth_base_url(self) -> str:
        tenant = quote(self.tenant, safe="")
        return f"{MICROSOFT_LOGIN_BASE_URL}/{tenant}/oauth2/v2.0"

    @staticmethod
    def _normalize_remote_folder(remote_folder: str) -> str:
        folder = remote_folder.replace("\\", "/").strip("/")
        parts = [part.strip() for part in folder.split("/") if part.strip()]
        return "/".join(parts)

    def authenticate(self) -> None:
        device = _post_form(
            f"{self.oauth_base_url}/devicecode",
            {
                "client_id": self.client_id,
                "scope": ONEDRIVE_SCOPES,
            },
        )
        verification_url = (
            device.get("verification_uri_complete")
            or device.get("verification_uri")
            or "https://microsoft.com/devicelogin"
        )
        print("\nOneDrive login required.")
        print(device.get("message") or f"Open {verification_url} and enter the code.")
        if self.open_browser:
            webbrowser.open(verification_url)

        interval = int(device.get("interval", 5))
        expires_at = time.monotonic() + int(device.get("expires_in", 900))
        token_url = f"{self.oauth_base_url}/token"
        while time.monotonic() < expires_at:
            time.sleep(interval)
            try:
                token = _post_form(
                    token_url,
                    {
                        "grant_type": "urn:ietf:params:oauth:grant-type:device_code",
                        "client_id": self.client_id,
                        "device_code": device["device_code"],
                    },
                )
            except OneDriveError as exc:
                error = exc.payload.get("error")
                if error == "authorization_pending":
                    continue
                if error == "slow_down":
                    interval += 5
                    continue
                if error == "expired_token":
                    raise OneDriveError("OneDrive login expired; run the command again.") from exc
                raise

            self._set_token(token)
            print("OneDrive login complete.")
            return

        raise OneDriveError("OneDrive login timed out; run the command again.")

    def _set_token(self, token: dict) -> None:
        self.access_token = token["access_token"]
        self.refresh_token = token.get("refresh_token", self.refresh_token)
        self.expires_at = time.time() + int(token.get("expires_in", 3600))

    def _ensure_token(self) -> None:
        if self.access_token and time.time() < self.expires_at - 120:
            return
        if not self.refresh_token:
            self.authenticate()
            return
        token = _post_form(
            f"{self.oauth_base_url}/token",
            {
                "grant_type": "refresh_token",
                "client_id": self.client_id,
                "refresh_token": self.refresh_token,
                "scope": ONEDRIVE_SCOPES,
            },
        )
        self._set_token(token)

    def _graph_json(
        self,
        method: str,
        path: str,
        *,
        body: dict | bytes | None = None,
        ok_statuses: set[int] | None = None,
        content_type: str = "application/json",
    ) -> dict:
        data = None
        if isinstance(body, dict):
            data = json.dumps(body).encode("utf-8")
        elif isinstance(body, bytes):
            data = body

        retry_statuses = {429, 500, 502, 503, 504}
        attempts = 8
        for attempt in range(1, attempts + 1):
            self._ensure_token()
            headers = {"Authorization": f"Bearer {self.access_token}"}
            if data is not None:
                headers["Content-Type"] = content_type
            req = Request(
                f"{GRAPH_BASE_URL}{path}",
                data=data,
                headers=headers,
                method=method,
            )
            try:
                return _request_json(req, ok_statuses=ok_statuses)
            except OneDriveError as exc:
                if exc.status == 401 and self.refresh_token:
                    self.access_token = ""
                    continue
                if exc.status in retry_statuses and attempt < attempts:
                    wait = exc.retry_after or min(120, 2 ** attempt)
                    print(
                        f"OneDrive is busy/throttling (HTTP {exc.status}); "
                        f"waiting {wait}s before retry {attempt + 1}/{attempts}..."
                    )
                    time.sleep(wait)
                    continue
                raise

        if self.refresh_token:
            try:
                self.access_token = ""
                self._ensure_token()
            except OneDriveError:
                pass
        raise OneDriveError("Microsoft Graph request failed after retries")

    def ensure_remote_folder(self) -> None:
        if self._folder_ready or not self.remote_folder:
            self._folder_ready = True
            return

        parent_parts: list[str] = []
        for part in self.remote_folder.split("/"):
            current_path = "/".join(parent_parts + [part])
            encoded_current = quote(current_path, safe="/")
            try:
                item = self._graph_json("GET", f"/me/drive/root:/{encoded_current}:")
                if "folder" not in item:
                    raise OneDriveError(f"OneDrive path is not a folder: {current_path}")
            except OneDriveError as exc:
                if exc.status != 404:
                    raise
                if parent_parts:
                    encoded_parent = quote("/".join(parent_parts), safe="/")
                    children_path = f"/me/drive/root:/{encoded_parent}:/children"
                else:
                    children_path = "/me/drive/root/children"
                self._graph_json(
                    "POST",
                    children_path,
                    body={
                        "name": part,
                        "folder": {},
                        "@microsoft.graph.conflictBehavior": "fail",
                    },
                )
            parent_parts.append(part)

        self._folder_ready = True

    def upload_file(self, local_path: Path) -> str:
        self.ensure_remote_folder()
        remote_path = local_path.name
        if self.remote_folder:
            remote_path = f"{self.remote_folder}/{remote_path}"
        encoded_remote = quote(remote_path, safe="/")
        with local_path.open("rb") as fh:
            body = fh.read()
        item = self._graph_json(
            "PUT",
            f"/me/drive/root:/{encoded_remote}:/content",
            body=body,
            ok_statuses={200, 201},
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        if self.upload_delay:
            time.sleep(self.upload_delay)
        return item.get("webUrl") or remote_path


def convert_tabular(
    input_path: Path,
    out_dir: Path,
    *,
    sheet_name: str | None = None,
    force: bool = False,
    limit: int | None = None,
    start_row: int = 2,
    uploader: OneDriveUploader | None = None,
    delete_after_upload: bool = False,
    resume_path: Path | None = None,
    debug: bool = False,
) -> ConversionSummary:
    out_dir.mkdir(parents=True, exist_ok=True)
    summary = ConversionSummary()
    considered = 0

    for row_number, row in iter_tabular_rows(input_path, sheet_name):
        if row_number < start_row:
            continue
        if limit is not None and considered >= limit:
            break
        considered += 1
        summary.total += 1

        out_path = tabular_output_path(row, row_number, out_dir)
        if out_path.exists() and not force:
            summary.skipped += 1
            print(f"Skip  {out_path.name} (already exists)")
            if resume_path:
                save_resume_row(resume_path, input_path, row_number + 1)
            continue

        try:
            blocks = tabular_row_to_blocks(
                row, row_number=row_number, source_name=input_path.name
            )
            if debug:
                print(f"\n=== {input_path.name} row {row_number} ===")
                for b in blocks:
                    print(f"  [{b.kind:10}] {b.text[:70]}")
                    for d in b.details:
                        print(f"               - {d[:66]}")
            _write_docx_atomic(blocks, out_path)
            if uploader:
                remote_url = uploader.upload_file(out_path)
                print(f"Uploaded {out_path.name} -> {remote_url}")
                if delete_after_upload:
                    out_path.unlink(missing_ok=True)
                else:
                    print(f"Wrote {out_path}")
            else:
                print(f"Wrote {out_path}")
            summary.converted += 1
            if resume_path:
                save_resume_row(resume_path, input_path, row_number + 1)
            if summary.total % 25 == 0:
                print(
                    f"Progress: {summary.converted} converted/uploaded, "
                    f"{summary.skipped} skipped, {summary.failed} failed "
                    f"({summary.total} rows checked)."
                )
        except Exception as exc:
            summary.failed += 1
            print(f"FAIL  {input_path.name} row {row_number}: {exc}")

    return summary


# --------------------------------------------------------------------------- #
# Orchestration
# --------------------------------------------------------------------------- #

def convert(
    image_path: Path,
    out_path: Path,
    *,
    scale: float = 1.5,
    cutoff_ratio: float = 0.66,
    keep_promo: bool = False,
    min_conf: int = 35,
    debug: bool = False,
) -> Path:
    img = load_image(image_path, scale)
    img, header_end = deinvert_dark_header(img)

    # First pass on the full image to locate the column divider.
    full_words = ocr_words(img, psm=3, min_conf=35)
    cutoff = find_column_cutoff(full_words, img.width, header_end, cutoff_ratio)

    # Crop to the main column and OCR as a single column. The confidence floor is
    # kept moderate: the entry logos OCR into low-confidence junk that is normal
    # height and tight to the text, so it can't be cleaned up downstream — letting
    # it in (min_conf 0) adds far more noise than the rare faint token it recovers.
    main = img.crop((0, 0, cutoff, img.height))
    words = ocr_words(main, psm=4, min_conf=min_conf)
    lines = group_lines(words, header_end)
    blocks = build_blocks(lines, header_end, keep_promo)

    if debug:
        print(f"\n=== {image_path.name} ===")
        print(f"header_end={header_end}  cutoff={cutoff}/{img.width}  "
              f"lines={len(lines)} blocks={len(blocks)}")
        for b in blocks:
            print(f"  [{b.kind:10}] {b.text[:70]}")
            for d in b.details:
                print(f"               - {d[:66]}")

    # Write to a temp file then atomically rename, so a run killed mid-write
    # never leaves a partial .docx that a later resume would wrongly skip.
    _write_docx_atomic(blocks, out_path)
    return out_path


def iter_inputs(input_path: Path):
    if input_path.is_dir():
        for p in sorted(input_path.iterdir()):
            if p.suffix.lower() in IMAGE_EXTS | TABULAR_EXTS:
                yield p
    else:
        yield input_path


def _is_image_input(path: Path) -> bool:
    return path.suffix.lower() in IMAGE_EXTS


def _is_tabular_input(path: Path) -> bool:
    return path.suffix.lower() in TABULAR_EXTS


def main(argv: list[str] | None = None) -> int:
    ap = argparse.ArgumentParser(
        description="Convert Doximity screenshots or Excel/CSV provider rows to Word docs."
    )
    ap.add_argument("input", help="Image, Excel/CSV file, or a folder of supported files")
    ap.add_argument("-o", "--output",
                    help="Output .docx for one image, output folder for "
                         "batches/Excel, or local staging folder with "
                         "--onedrive-upload")
    ap.add_argument("--tesseract", help="Path to tesseract.exe")
    ap.add_argument("--scale", type=float, default=1.5,
                    help="Upscale factor before OCR (default 1.5)")
    ap.add_argument("--cutoff-ratio", type=float, default=0.66,
                    help="Fallback main-column width fraction (default 0.66)")
    ap.add_argument("--min-conf", type=int, default=35,
                    help="Min OCR confidence for body text (default 35; lower "
                         "toward 0 for max recall, but expect icon/seal noise)")
    ap.add_argument("--keep-promo", action="store_true",
                    help="Keep Doximity promo/navigation lines")
    ap.add_argument("--force", action="store_true",
                    help="Re-process rows/images even if their .docx already exists "
                         "(default: skip existing docs so a stopped run resumes)")
    ap.add_argument("--sheet",
                    help="Worksheet name for Excel inputs (default: active worksheet)")
    ap.add_argument("--limit", type=int,
                    help="Maximum number of Excel/CSV data rows to process")
    ap.add_argument("--start-row", type=int,
                    help="First Excel/CSV data row number to process. Overrides "
                         "the saved resume checkpoint")
    ap.add_argument("--no-resume", action="store_true",
                    help="Ignore the saved Excel/CSV row checkpoint and start "
                         "from --start-row or row 2")
    ap.add_argument("--onedrive-upload", action="store_true",
                    help="Upload generated .docx files to OneDrive using a "
                         "browser/device-code Microsoft login")
    ap.add_argument("--onedrive-client-id",
                    default=os.environ.get("ONEDRIVE_CLIENT_ID"),
                    help="Microsoft Graph public client ID. Can also be set "
                         "with the ONEDRIVE_CLIENT_ID environment variable")
    ap.add_argument("--onedrive-tenant",
                    default=os.environ.get("ONEDRIVE_TENANT", "consumers"),
                    help="Microsoft login tenant/authority for OneDrive auth: "
                         "consumers for personal accounts, common for any "
                         "configured account type, or a directory tenant ID "
                         "(default: consumers)")
    ap.add_argument("--onedrive-folder", default="Converter Output",
                    help="Remote OneDrive folder under root (default: "
                         "'Converter Output')")
    ap.add_argument("--onedrive-upload-delay", type=float, default=0.0,
                    help="Seconds to wait after each OneDrive upload. Use a "
                         "small value like 0.5 if your tenant throttles long runs")
    ap.add_argument("--keep-local", action="store_true",
                    help="With --onedrive-upload, keep the local .docx files "
                         "after upload")
    ap.add_argument("--no-browser", action="store_true",
                    help="With --onedrive-upload, print the login URL/code "
                         "without trying to open a browser")
    ap.add_argument("--debug", action="store_true",
                    help="Print the detected structure")
    args = ap.parse_args(argv)

    input_path = Path(args.input)
    if not input_path.exists():
        sys.exit(f"Input not found: {input_path}")

    inputs = list(iter_inputs(input_path))
    if not inputs:
        sys.exit(f"No supported image, Excel, or CSV files found in {input_path}")

    unsupported = [
        p for p in inputs
        if not _is_image_input(p) and not _is_tabular_input(p)
    ]
    if unsupported:
        supported = ", ".join(sorted(IMAGE_EXTS | TABULAR_EXTS))
        sys.exit(f"Unsupported input type: {unsupported[0]} (supported: {supported})")

    if args.limit is not None and args.limit < 1:
        sys.exit("--limit must be 1 or greater")
    if args.start_row is not None and args.start_row < 2:
        sys.exit("--start-row must be 2 or greater")
    if args.onedrive_upload_delay < 0:
        sys.exit("--onedrive-upload-delay must be 0 or greater")
    if args.onedrive_upload and not args.onedrive_client_id:
        sys.exit(
            "--onedrive-upload requires --onedrive-client-id or the "
            "ONEDRIVE_CLIENT_ID environment variable."
        )

    if any(_is_image_input(p) for p in inputs):
        configure_tesseract(args.tesseract)

    uploader = None
    if args.onedrive_upload:
        uploader = OneDriveUploader(
            client_id=args.onedrive_client_id,
            tenant=args.onedrive_tenant,
            remote_folder=args.onedrive_folder,
            upload_delay=args.onedrive_upload_delay,
            open_browser=not args.no_browser,
        )
        try:
            uploader.authenticate()
        except OneDriveError as exc:
            hint = ""
            if "AADSTS50059" in str(exc):
                hint = (
                    "\nHint: Try adding `--onedrive-tenant consumers` for a "
                    "personal OneDrive account, or use your Microsoft Entra "
                    "Directory (tenant) ID for a work/school account. Also "
                    "check that the app registration allows public client flows."
                )
            sys.exit(f"OneDrive login failed: {exc}{hint}")

    out_arg = Path(args.output) if args.output else None
    has_tabular = any(_is_tabular_input(p) for p in inputs)
    out_dir = None
    temp_dir = None
    if has_tabular and out_arg and out_arg.suffix.lower() == ".docx":
        sys.exit("Excel/CSV inputs write one .docx per row; pass an output folder.")

    if args.onedrive_upload and not out_arg:
        temp_dir = tempfile.TemporaryDirectory(prefix="converter_onedrive_")
        out_dir = Path(temp_dir.name)
        out_dir.mkdir(parents=True, exist_ok=True)
    elif has_tabular:
        if input_path.is_dir() or len(inputs) > 1:
            out_dir = out_arg or (input_path if input_path.is_dir() else input_path.parent)
        else:
            out_dir = out_arg or input_path.with_name(f"{input_path.stem}_docs")
        out_dir.mkdir(parents=True, exist_ok=True)
    elif input_path.is_dir() or len(inputs) > 1:
        out_dir = out_arg or (input_path if input_path.is_dir() else input_path.parent)
        out_dir.mkdir(parents=True, exist_ok=True)

    summary = ConversionSummary()
    for source_path in inputs:
        if _is_tabular_input(source_path):
            source_start_row = args.start_row or 2
            state_path = None
            if not args.no_resume:
                if uploader:
                    output_identity = (
                        f"onedrive:{args.onedrive_tenant}:"
                        f"{args.onedrive_folder}:{args.onedrive_client_id}"
                    )
                else:
                    output_identity = f"local:{out_dir}"
                state_path = resume_state_path(
                    source_path,
                    output_identity=output_identity,
                    sheet_name=args.sheet,
                )
                if args.start_row is None:
                    saved_row = load_resume_row(state_path, source_path)
                    if saved_row and saved_row > 2:
                        source_start_row = saved_row
                        print(
                            f"Resuming {source_path.name} from Excel row "
                            f"{source_start_row} using {state_path}"
                        )

            try:
                tabular_summary = convert_tabular(
                    source_path,
                    out_dir,
                    sheet_name=args.sheet,
                    force=args.force,
                    limit=args.limit,
                    start_row=source_start_row,
                    uploader=uploader,
                    delete_after_upload=bool(uploader and not args.keep_local),
                    resume_path=state_path,
                    debug=args.debug,
                )
                summary.converted += tabular_summary.converted
                summary.skipped += tabular_summary.skipped
                summary.failed += tabular_summary.failed
                summary.total += tabular_summary.total
            except Exception as exc:
                summary.failed += 1
                summary.total += 1
                print(f"FAIL  {source_path.name}: {exc}")
            continue

        if out_dir is not None:
            out_path = out_dir / (source_path.stem + ".docx")
        else:
            out_path = out_arg or source_path.with_suffix(".docx")

        # Resume support: a finished run leaves a .docx per image, so skip any
        # that already exist unless --force is given.
        if out_path.exists() and not args.force:
            summary.skipped += 1
            summary.total += 1
            print(f"Skip  {out_path.name} (already exists)")
            continue

        try:
            convert(
                source_path, out_path,
                scale=args.scale, cutoff_ratio=args.cutoff_ratio,
                keep_promo=args.keep_promo, min_conf=args.min_conf, debug=args.debug,
            )
            if uploader:
                remote_url = uploader.upload_file(out_path)
                print(f"Uploaded {out_path.name} -> {remote_url}")
                if not args.keep_local:
                    out_path.unlink(missing_ok=True)
                else:
                    print(f"Wrote {out_path}")
            else:
                print(f"Wrote {out_path}")
            summary.converted += 1
            summary.total += 1
        except Exception as exc:  # keep going through the batch on a single failure
            summary.failed += 1
            summary.total += 1
            print(f"FAIL  {source_path.name}: {exc}")

    print(f"\nDone. {summary.converted} converted, {summary.skipped} skipped, "
          f"{summary.failed} failed ({summary.total} total).")
    if temp_dir is not None:
        temp_dir.cleanup()
    return 1 if summary.failed else 0


if __name__ == "__main__":
    raise SystemExit(main())
